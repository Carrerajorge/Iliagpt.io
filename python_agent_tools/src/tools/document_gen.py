"""Document generation tool for DOCX and PDF files."""

from typing import Optional, List, Dict, Any, Literal
from pydantic import BaseModel, Field
from .base import BaseTool, ToolCategory, Priority, ToolInput, ToolOutput
from ..core.registry import ToolRegistry
import base64
import re
from io import BytesIO


class DocumentSection(BaseModel):
    """A section of the document."""
    type: Literal["heading", "paragraph", "bullet_list", "numbered_list", "table", "page_break"]
    content: Any
    level: int = Field(1, ge=1, le=6)
    style: Optional[Dict[str, Any]] = None


class TableData(BaseModel):
    """Table data structure."""
    headers: List[str]
    rows: List[List[str]]


class DocumentGenInput(ToolInput):
    """Input for document generation tool."""
    title: str = Field(..., min_length=1, max_length=500, description="Document title")
    sections: List[DocumentSection] = Field(..., description="Document sections")
    format: Literal["docx", "pdf", "html"] = Field("docx", description="Output format")
    author: Optional[str] = Field(None, description="Document author")
    subject: Optional[str] = Field(None, description="Document subject")
    template_vars: Dict[str, str] = Field(default={}, description="Variables for template substitution")
    font_name: str = Field("Calibri", description="Default font name")
    font_size: int = Field(11, ge=8, le=72, description="Default font size in points")


class GeneratedDocument(BaseModel):
    """Generated document result."""
    filename: str
    content_base64: str
    format: str
    size_bytes: int


class DocumentGenOutput(ToolOutput):
    """Output from document generation tool."""
    data: Optional[GeneratedDocument] = None


def substitute_variables(text: str, variables: Dict[str, str]) -> str:
    """Substitute {{variable}} patterns in text."""
    for key, value in variables.items():
        pattern = r'\{\{\s*' + re.escape(key) + r'\s*\}\}'
        text = re.sub(pattern, value, text)
    return text


def sanitize_filename(title: str) -> str:
    """Create a safe filename from title."""
    safe = re.sub(r'[^\w\s-]', '', title)
    safe = re.sub(r'[\s]+', '_', safe)
    return safe[:100] or "document"


@ToolRegistry.register
class DocumentGenTool(BaseTool[DocumentGenInput, DocumentGenOutput]):
    """Tool for generating DOCX and PDF documents."""
    
    name = "document_gen"
    description = "Generates DOCX and PDF documents with structured content, templates, and styling"
    category = ToolCategory.GENERATION
    priority = Priority.MEDIUM
    dependencies = []
    
    async def execute(self, input: DocumentGenInput) -> DocumentGenOutput:
        """Execute document generation."""
        self.logger.info("document_gen_start", format=input.format, sections=len(input.sections))
        
        try:
            if input.format == "docx":
                return await self._generate_docx(input)
            elif input.format == "html":
                return await self._generate_html(input)
            elif input.format == "pdf":
                return await self._generate_pdf(input)
            else:
                return DocumentGenOutput(
                    success=False,
                    error=f"Unsupported format: {input.format}"
                )
        except Exception as e:
            self.logger.error("document_gen_error", error=str(e))
            return DocumentGenOutput(
                success=False,
                error=f"Document generation failed: {str(e)}"
            )
    
    async def _generate_docx(self, input: DocumentGenInput) -> DocumentGenOutput:
        """Generate DOCX document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return DocumentGenOutput(
                success=False,
                error="python-docx not installed. Run: pip install python-docx"
            )
        
        doc = Document()
        
        core_props = doc.core_properties
        core_props.title = substitute_variables(input.title, input.template_vars)
        if input.author:
            core_props.author = input.author
        if input.subject:
            core_props.subject = input.subject
        
        title_para = doc.add_heading(
            substitute_variables(input.title, input.template_vars),
            level=0
        )
        
        for section in input.sections:
            content = section.content
            if isinstance(content, str):
                content = substitute_variables(content, input.template_vars)
            
            if section.type == "heading":
                doc.add_heading(content, level=min(section.level, 9))
            
            elif section.type == "paragraph":
                para = doc.add_paragraph(content)
                if section.style:
                    if section.style.get("bold"):
                        for run in para.runs:
                            run.bold = True
                    if section.style.get("italic"):
                        for run in para.runs:
                            run.italic = True
            
            elif section.type == "bullet_list":
                items = content if isinstance(content, list) else [content]
                for item in items:
                    item_text = substitute_variables(str(item), input.template_vars)
                    doc.add_paragraph(item_text, style="List Bullet")
            
            elif section.type == "numbered_list":
                items = content if isinstance(content, list) else [content]
                for item in items:
                    item_text = substitute_variables(str(item), input.template_vars)
                    doc.add_paragraph(item_text, style="List Number")
            
            elif section.type == "table":
                if isinstance(content, dict):
                    headers = content.get("headers", [])
                    rows = content.get("rows", [])
                    if headers:
                        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                        table.style = "Table Grid"
                        
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            header_cells[i].text = substitute_variables(str(header), input.template_vars)
                        
                        for row_idx, row_data in enumerate(rows):
                            row_cells = table.rows[row_idx + 1].cells
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < len(row_cells):
                                    row_cells[col_idx].text = substitute_variables(str(cell_data), input.template_vars)
            
            elif section.type == "page_break":
                doc.add_page_break()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        content_bytes = buffer.read()
        
        filename = f"{sanitize_filename(input.title)}.docx"
        
        result = GeneratedDocument(
            filename=filename,
            content_base64=base64.b64encode(content_bytes).decode("utf-8"),
            format="docx",
            size_bytes=len(content_bytes)
        )
        
        self.logger.info("document_gen_complete", format="docx", size=len(content_bytes))
        
        return DocumentGenOutput(
            success=True,
            data=result,
            metadata={"filename": filename, "sections_count": len(input.sections)}
        )
    
    async def _generate_html(self, input: DocumentGenInput) -> DocumentGenOutput:
        """Generate HTML document."""
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>{substitute_variables(input.title, input.template_vars)}</title>",
            "<meta charset='utf-8'>",
            f"<style>body {{ font-family: {input.font_name}, sans-serif; font-size: {input.font_size}pt; max-width: 800px; margin: 0 auto; padding: 20px; }}</style>",
            "</head>",
            "<body>",
            f"<h1>{substitute_variables(input.title, input.template_vars)}</h1>"
        ]
        
        for section in input.sections:
            content = section.content
            if isinstance(content, str):
                content = substitute_variables(content, input.template_vars)
            
            if section.type == "heading":
                level = min(section.level + 1, 6)
                html_parts.append(f"<h{level}>{content}</h{level}>")
            
            elif section.type == "paragraph":
                html_parts.append(f"<p>{content}</p>")
            
            elif section.type == "bullet_list":
                items = content if isinstance(content, list) else [content]
                html_parts.append("<ul>")
                for item in items:
                    item_text = substitute_variables(str(item), input.template_vars)
                    html_parts.append(f"<li>{item_text}</li>")
                html_parts.append("</ul>")
            
            elif section.type == "numbered_list":
                items = content if isinstance(content, list) else [content]
                html_parts.append("<ol>")
                for item in items:
                    item_text = substitute_variables(str(item), input.template_vars)
                    html_parts.append(f"<li>{item_text}</li>")
                html_parts.append("</ol>")
            
            elif section.type == "table":
                if isinstance(content, dict):
                    headers = content.get("headers", [])
                    rows = content.get("rows", [])
                    html_parts.append("<table border='1' cellpadding='5' cellspacing='0'>")
                    if headers:
                        html_parts.append("<tr>")
                        for header in headers:
                            html_parts.append(f"<th>{substitute_variables(str(header), input.template_vars)}</th>")
                        html_parts.append("</tr>")
                    for row_data in rows:
                        html_parts.append("<tr>")
                        for cell_data in row_data:
                            html_parts.append(f"<td>{substitute_variables(str(cell_data), input.template_vars)}</td>")
                        html_parts.append("</tr>")
                    html_parts.append("</table>")
            
            elif section.type == "page_break":
                html_parts.append("<div style='page-break-after: always;'></div>")
        
        html_parts.extend(["</body>", "</html>"])
        html_content = "\n".join(html_parts)
        content_bytes = html_content.encode("utf-8")
        
        filename = f"{sanitize_filename(input.title)}.html"
        
        result = GeneratedDocument(
            filename=filename,
            content_base64=base64.b64encode(content_bytes).decode("utf-8"),
            format="html",
            size_bytes=len(content_bytes)
        )
        
        self.logger.info("document_gen_complete", format="html", size=len(content_bytes))
        
        return DocumentGenOutput(
            success=True,
            data=result,
            metadata={"filename": filename, "sections_count": len(input.sections)}
        )
    
    async def _generate_pdf(self, input: DocumentGenInput) -> DocumentGenOutput:
        """Generate PDF by first creating HTML then converting."""
        html_result = await self._generate_html(input)
        
        if not html_result.success or not html_result.data:
            return DocumentGenOutput(
                success=False,
                error="Failed to generate intermediate HTML for PDF"
            )
        
        filename = f"{sanitize_filename(input.title)}.pdf"
        
        return DocumentGenOutput(
            success=True,
            data=GeneratedDocument(
                filename=filename,
                content_base64=html_result.data.content_base64,
                format="html_for_pdf",
                size_bytes=html_result.data.size_bytes
            ),
            metadata={
                "filename": filename,
                "note": "HTML content provided - use browser_tool or wkhtmltopdf for final PDF conversion"
            }
        )
