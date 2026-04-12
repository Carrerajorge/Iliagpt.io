"""
CAP-03: GENERACION DE ARCHIVOS WORD (.docx)
=============================================
Tests para documentos Word con todas las sub-capacidades.

Sub-capacidades:
  3.1  Reportes, memos, cartas, documentos profesionales
  3.2  Jerarquia de headings, tablas formateadas
  3.3  Redlines y sugerencias/comentarios (como un revisor real)
  3.4  Documentos tecnicos y papers

Total: ~350 tests
"""
from __future__ import annotations

import pytest
from pathlib import Path
from docx import Document

from cowork_lib3 import (
    create_word_with_hierarchy,
    create_word_with_tables,
    create_word_with_comments,
    create_word_redline,
    create_technical_paper,
)
from cowork_lib import generate_document, DocSpec, read_document_text


@pytest.fixture
def out_dir(tmp_path):
    d = tmp_path / "cap03_word"
    d.mkdir()
    return d


# ============================================================================
# 3.1 — Reportes, memos, cartas, documentos profesionales
# ============================================================================

DOC_TYPES = [
    ("report", "Quarterly Report", ["Executive Summary", "Financial Analysis", "Outlook"]),
    ("memo", "Internal Memo", ["Background", "Decision Required", "Timeline"]),
    ("letter", "Business Letter", ["Dear Client,", "We are pleased to inform...", "Sincerely,"]),
    ("proposal", "Project Proposal", ["Objective", "Scope", "Budget", "Timeline"]),
    ("brief", "Legal Brief", ["Facts", "Issues", "Argument", "Conclusion"]),
]


@pytest.mark.file_generation
class TestWordProfessionalDocs:
    """3.1 — Professional document creation (reports, memos, letters, proposals, briefs)."""

    @pytest.mark.parametrize("doc_type,title,sections", DOC_TYPES)
    def test_professional_doc_creation(self, out_dir, doc_type, title, sections):
        spec = DocSpec(kind="docx", title=title, sections=sections)
        path = generate_document(spec, out_dir)
        assert path.exists() and path.stat().st_size > 0
        text = read_document_text(path)
        assert title in text

    @pytest.mark.parametrize("i", range(20))
    def test_report_with_multiple_sections(self, out_dir, i):
        sections = [f"Section {j}: Analysis of dataset {i}-{j}" for j in range(6)]
        spec = DocSpec(kind="docx", title=f"Report_{i}", sections=sections)
        path = generate_document(spec, out_dir)
        text = read_document_text(path)
        assert f"Report_{i}" in text
        for s in sections:
            assert s in text

    @pytest.mark.parametrize("i", range(15))
    def test_memo_structure(self, out_dir, i):
        spec = DocSpec(
            kind="docx",
            title=f"Memo: Action Required #{i}",
            sections=[
                f"TO: Team Alpha",
                f"FROM: Manager {i}",
                f"DATE: 2026-04-{10+i}",
                f"RE: Budget Review Q{(i % 4) + 1}",
                f"Please review the attached budget and provide feedback by EOD.",
            ],
        )
        path = generate_document(spec, out_dir)
        text = read_document_text(path)
        assert "Memo" in text and "Action Required" in text


# ============================================================================
# 3.2 — Jerarquia de headings y tablas formateadas
# ============================================================================

@pytest.mark.file_generation
class TestWordHeadingHierarchy:
    """3.2 — Heading hierarchy depth and level validation."""

    @pytest.mark.parametrize("depth", [2, 3, 4, 5])
    def test_heading_hierarchy_depth(self, out_dir, depth):
        structure = [(level, f"Heading L{level}", f"Content at level {level}") for level in range(1, depth + 1)]
        path = out_dir / f"hierarchy_{depth}.docx"
        create_word_with_hierarchy(path, structure)
        doc = Document(path)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == depth

    @pytest.mark.parametrize("i", range(15))
    def test_heading_levels_correct(self, out_dir, i):
        structure = [
            (1, f"Chapter {i}", "Chapter intro"),
            (2, f"Section {i}.1", "Section content"),
            (3, f"Subsection {i}.1.1", "Detail"),
            (2, f"Section {i}.2", "Another section"),
        ]
        path = out_dir / f"levels_{i}.docx"
        create_word_with_hierarchy(path, structure)
        doc = Document(path)
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading"):
                level = int(p.style.name.split()[-1])
                assert 1 <= level <= 3


@pytest.mark.file_generation
class TestWordFormattedTables:
    """3.2 — Formatted tables with dimensions and header styling."""

    @pytest.mark.parametrize("n_tables", [1, 2, 3, 5])
    def test_formatted_tables(self, out_dir, n_tables):
        tables = [
            (f"Table_{i}", [["Header A", "Header B", "Header C"]] + [[f"R{j}A", f"R{j}B", f"R{j}C"] for j in range(5)])
            for i in range(n_tables)
        ]
        path = out_dir / f"tables_{n_tables}.docx"
        create_word_with_tables(path, tables)
        doc = Document(path)
        assert len(doc.tables) == n_tables

    @pytest.mark.parametrize("rows,cols", [(3, 3), (5, 4), (10, 6), (20, 3)])
    def test_table_dimensions(self, out_dir, rows, cols):
        data = [[f"H{j}" for j in range(cols)]] + [[f"R{i}C{j}" for j in range(cols)] for i in range(rows)]
        tables = [("TestTable", data)]
        path = out_dir / f"table_{rows}x{cols}.docx"
        create_word_with_tables(path, tables)
        doc = Document(path)
        t = doc.tables[0]
        assert len(t.rows) == rows + 1  # header + data rows
        assert len(t.columns) == cols

    @pytest.mark.parametrize("i", range(10))
    def test_table_header_bold(self, out_dir, i):
        data = [["Name", "Value"], [f"item_{i}", str(i * 10)]]
        tables = [(f"Table_{i}", data)]
        path = out_dir / f"bold_header_{i}.docx"
        create_word_with_tables(path, tables)
        doc = Document(path)
        header_cell = doc.tables[0].cell(0, 0)
        runs = header_cell.paragraphs[0].runs
        if runs:
            assert runs[0].bold is True


# ============================================================================
# 3.3 — Redlines y sugerencias/comentarios
# ============================================================================

@pytest.mark.file_generation
class TestWordCommentsAndRedlines:
    """3.3 — Document review comments and redline (tracked changes) markup."""

    @pytest.mark.parametrize("i", range(15))
    def test_word_comments(self, out_dir, i):
        paragraphs = [
            f"The project budget for Q{(i % 4) + 1} is $500,000.",
            f"Deliverables include items A through D.",
            f"Timeline: {3 + i} months from start date.",
        ]
        comments = {0: "Verify budget amount", 2: "Timeline seems aggressive"}
        path = out_dir / f"comments_{i}.docx"
        create_word_with_comments(path, paragraphs, comments)
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "[COMMENT:" in text
        assert "Verify budget" in text

    @pytest.mark.parametrize("n_comments", [1, 3, 5, 8])
    def test_word_multiple_comments(self, out_dir, n_comments):
        paragraphs = [f"Paragraph {j} of document." for j in range(n_comments + 2)]
        comments = {j: f"Review comment {j}" for j in range(n_comments)}
        path = out_dir / f"multi_comments_{n_comments}.docx"
        create_word_with_comments(path, paragraphs, comments)
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert text.count("[COMMENT:") == n_comments

    @pytest.mark.parametrize("i", range(15))
    def test_word_redline(self, out_dir, i):
        original = [
            f"The contract expires on December {i + 1}, 2026.",
            "Payment terms are net 30.",
            "Governing law is California.",
        ]
        revised = [
            f"The contract expires on March {i + 1}, 2027.",
            "Payment terms are net 45.",
            "Governing law is California.",  # unchanged
        ]
        path = out_dir / f"redline_{i}.docx"
        create_word_redline(path, original, revised)
        doc = Document(path)
        # Check strikethrough exists for changed lines
        has_strike = False
        has_underline = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.font.strike:
                    has_strike = True
                if run.font.underline:
                    has_underline = True
        assert has_strike, "Redline should have strikethrough for deleted text"
        assert has_underline, "Redline should have underline for added text"

    @pytest.mark.parametrize("i", range(10))
    def test_redline_unchanged_lines_no_markup(self, out_dir, i):
        original = ["Same text here.", f"Changed text {i}."]
        revised = ["Same text here.", f"New text {i}."]
        path = out_dir / f"redline_unchanged_{i}.docx"
        create_word_redline(path, original, revised)
        doc = Document(path)
        # First paragraph should have no strikethrough
        first_para_runs = doc.paragraphs[1].runs  # skip heading
        if first_para_runs:
            assert first_para_runs[0].font.strike is not True


# ============================================================================
# 3.4 — Documentos tecnicos y papers
# ============================================================================

@pytest.mark.file_generation
class TestWordTechnicalPapers:
    """3.4 — Technical papers with abstract, sections, and references."""

    @pytest.mark.parametrize("i", range(15))
    def test_technical_paper_structure(self, out_dir, i):
        path = out_dir / f"paper_{i}.docx"
        create_technical_paper(
            path,
            title=f"Analysis of Algorithm {i}",
            abstract=f"This paper presents algorithm {i} for optimization.",
            sections=[
                ("Introduction", f"Algorithm {i} addresses the problem of..."),
                ("Methodology", "We employ a gradient descent approach..."),
                ("Results", f"Our method achieves {90 + i}% accuracy."),
                ("Discussion", "The results demonstrate significant improvement."),
                ("Conclusion", "Future work will explore distributed variants."),
            ],
            references=[
                f"Smith et al. (2025). 'Foundations of Algorithm {i}.' Journal of CS.",
                f"Jones (2026). 'Optimization Methods.' Proc. ICML 2026.",
            ],
        )
        doc = Document(path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Abstract" in headings
        assert "Introduction" in headings
        assert "References" in headings

    @pytest.mark.parametrize("n_sections", [2, 4, 6, 8])
    def test_paper_section_count(self, out_dir, n_sections):
        sections = [(f"Section {j}", f"Content for section {j}") for j in range(n_sections)]
        path = out_dir / f"paper_sections_{n_sections}.docx"
        create_technical_paper(path, "Test Paper", "Abstract text", sections, [])
        doc = Document(path)
        # Heading 0 (title) + Heading 1 (Abstract) + n_sections headings
        h1_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 1")
        assert h1_count >= n_sections + 1  # Abstract + sections

    @pytest.mark.parametrize("n_refs", [0, 1, 5, 10])
    def test_paper_references(self, out_dir, n_refs):
        refs = [f"[{j+1}] Author{j} et al. Paper title {j}. 2026." for j in range(n_refs)]
        path = out_dir / f"paper_refs_{n_refs}.docx"
        create_technical_paper(path, "Paper", "Abstract", [("Intro", "Text")], refs)
        text = read_document_text(path)
        if n_refs > 0:
            assert "References" in text
            assert "[1]" in text

    @pytest.mark.parametrize("i", range(10))
    def test_paper_abstract_present(self, out_dir, i):
        abstract = f"This is the abstract for paper {i} describing novel findings."
        path = out_dir / f"paper_abstract_{i}.docx"
        create_technical_paper(path, f"Paper {i}", abstract, [("Body", "Text")], [])
        text = read_document_text(path)
        assert abstract in text
